from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path(r"D:\Downloads-2\School\Cloud\DoAnCloud.docx")
OUT = Path(r"D:\Downloads-2\School\Other\websocial\DoAnCloud_cap_nhat.docx")
DIAGRAM = Path(r"D:\Downloads-2\School\Other\websocial\assets\luong-hoat-dong-aws-moi.png")

doc = Document(SRC)

def set_text(p, text):
    p.clear()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p

replacements = {
    "XÂY DỰNG BLOG CÁ NHÂN": "XÂY DỰNG VÀ TRIỂN KHAI MẠNG XÃ HỘI TRÊN AWS",
    "Xây dựng một trang blog cá nhân với đầy đủ các tính năng quản trị nội dung (CRUD), phân quyền và xác thực là một đề tài cơ bản và phổ biến trong môi trường học tập. Tuy nhiên, khoảng cách từ việc chạy ứng dụng thành công trên localhost đến việc triển khai thực tế trên nền tảng điện toán đám mây là một bước tiến đối với sinh viên. Đặc biệt, với kiến trúc sử dụng web chúng em sử dụng là Java/Spring Boot framework đòi hỏi nhiều tài nguyên bộ nhớ kết hợp cùng cơ sở dữ liệu MySQL, việc triển khai lên các dịch vụ đám mây thường gặp khó khăn về cạn kiệt bộ nhớ. Do đó, việc tìm ra giải pháp đưa hệ thống này lên môi trường mạng một cách ổn định với chi phí thấp nhất là một bài toán cấp thiết.": "Đề tài xây dựng một ứng dụng mạng xã hội bằng Java Spring Boot, hỗ trợ tài khoản người dùng, bài viết, bình luận, lượt thích, kết bạn, story, nhắn tin thời gian thực và khu vực quản trị. Trọng tâm của đồ án không chỉ là phát triển chức năng mà còn là thiết kế quy trình triển khai tự động trên AWS. Kiến trúc mới tách cơ sở dữ liệu khỏi máy chủ ứng dụng, sử dụng dịch vụ lưu trữ đối tượng cho tệp và áp dụng chuỗi CI/CD AWS-native nhằm tăng tính ổn định, khả năng bảo trì và mức độ sát với môi trường thực tế.",
    "Phân hệ giao diện người dùng: Xây dựng một trang hiển thị danh sách các bài viết cá nhân cho người dùng truy cập, được thiết kế tối giản, tập trung vào chức năng.": "Phân hệ người dùng: đăng ký, đăng nhập, hồ sơ cá nhân, bảng tin, khám phá nội dung, đăng bài, bình luận, thích, chia sẻ, kết bạn, story và nhắn tin thời gian thực qua WebSocket.",
    "Phân hệ quản trị: Tích hợp một khu vực điều khiển dành riêng cho quản trị viên, được bảo vệ nghiêm ngặt bằng cơ chế xác thực phiên đăng nhập. Người quản trị có toàn quyền thực hiện các thao tác cơ sở dữ liệu nền tảng bao gồm: Thêm bài viết mới, Xem chi tiết nội dung, Chỉnh sửa thông tin cá nhân và Xóa các bài viết.": "Phân hệ quản trị: cung cấp giao diện dành cho quản trị viên để theo dõi và quản lý người dùng, bài viết và nội dung vi phạm; quyền truy cập được kiểm soát bằng cơ chế xác thực của ứng dụng.",
    "Hạ tầng kỹ thuật triển khai: Không sử dụng hay xây dựng máy chủ vật lý độc lập. Toàn bộ mã nguồn Java/Spring và cơ sở dữ liệu MySQL sẽ được triển khai trực tiếp trên máy chủ ảo của nền tảng điện toán đám mây Amazon Web Services (AWS) kết hợp cùng tên miền động từ DuckDNS. Nginx làm reverse proxy và GitLab CI/CD để tự động hóa quá trình cập nhật phiên bản.": "Hạ tầng triển khai: ứng dụng Spring Boot chạy trên Amazon EC2; dữ liệu quan hệ được tách sang Amazon RDS for MySQL; tệp tải lên được định hướng lưu trên Amazon S3. Quy trình giao hàng phần mềm sử dụng AWS CodePipeline để điều phối, AWS CodeBuild để đóng gói Maven và AWS CodeDeploy để đưa bản phát hành lên EC2 theo tệp appspec.yml.",
    "Mô hình IaaS (Infrastructure as a Service): Cách thức hệ thống đám mây cấp phát tài nguyên máy chủ ảo (EC2), quản lý bộ nhớ, IP và tường lửa.": "Mô hình kết hợp IaaS và dịch vụ được quản lý: EC2 cung cấp máy chủ ứng dụng; RDS quản lý MySQL; S3 cung cấp lưu trữ đối tượng; Security Group kiểm soát luồng mạng.",
    "Kiến trúc Reverse Proxy: Nghiên cứu và triển khai Nginx làm máy chủ trung gian bảo vệ và điều hướng dữ liệu cho Spring Boot.": "Kiến trúc ứng dụng nhiều tầng: trình duyệt gửi yêu cầu đến Spring Boot trên EC2; ứng dụng truy cập RDS qua kết nối riêng và S3 qua API AWS.",
    "Tự động hóa hệ thống Linux: Sử dụng systemd và crontab để biến các tiến trình phần mềm thành các dịch vụ tự động trên Ubuntu và GitLab CI/CD để tự động build/deploy.": "Tự động hóa triển khai: CodePipeline nhận thay đổi mã nguồn, CodeBuild chạy Maven, CodeDeploy thực thi các lifecycle hook và systemd quản lý dịch vụ tc-social trên EC2.",
    "2.4. Nginx Reverse Proxy": "2.4. Các dịch vụ AWS sử dụng trong kiến trúc mới",
    "Web Server (Nginx Reverse Proxy):cơ chế đặt một máy chủ trung gian đứng trước ứng dụng backend. Trong đề tài, Nginx nhận request từ Internet ở cổng 80/443, sau đó chuyển tiếp vào Spring Boot đang chạy ở localhost:8080. Cách triển khai này giúp người dùng không truy cập trực tiếp vào cổng 8080, đồng thời hỗ trợ cấu hình HTTPS, chuyển hướng HTTP sang HTTPS và giới hạn dung lượng upload.": "Amazon EC2 vận hành ứng dụng Spring Boot; Amazon RDS for MySQL cung cấp cơ sở dữ liệu được quản lý và tách khỏi vòng đời máy chủ ứng dụng; Amazon S3 lưu trữ tệp theo mô hình đối tượng. Các Security Group chỉ cho phép những luồng kết nối cần thiết giữa người dùng, EC2 và RDS.",
    "2.5. Dynamic DNS với DuckDNS": "2.5. AWS CodePipeline, CodeBuild và CodeDeploy",
    "Khi sử dụng EC2 không gắn Elastic IP, địa chỉ public IPv4 có thể thay đổi sau mỗi lần stop/start instance. DuckDNS được sử dụng để cập nhật IP động cho tên miền tcblog-cloudcomputing.duckdns.org. Trên máy chủ EC2, crontab được cấu hình để gọi URL cập nhật IP định kỳ mỗi 5 phút, giúp tên miền luôn trỏ về public IP hiện tại của EC2.": "CodePipeline là lớp điều phối toàn bộ quy trình CI/CD. Khi kho mã có phiên bản mới, CodeBuild đọc buildspec.yml, sử dụng Amazon Corretto 21 và Maven để tạo tệp JAR cùng bộ tệp triển khai. CodeDeploy đọc appspec.yml, chép gói phát hành vào /opt/tc-social và chạy các hook dừng, cài đặt, khởi động ứng dụng.",
    "2.6. GitLab CI/CD trong triển khai phần mềm": "2.6. Quản lý cấu hình và dịch vụ ứng dụng",
    "CI/CD là quy trình tự động hóa các bước tích hợp và triển khai phần mềm. Trong đề tài, GitLab CI/CD được sử dụng để build project Spring Boot bằng Maven, tạo file JAR, truyền file JAR lên EC2 thông qua SSH/SCP và khởi động lại systemd service. Nhờ đó, khi lập trình viên push code lên nhánh main, hệ thống có thể tự động cập nhật phiên bản mới mà không cần thao tác upload file thủ công.": "Hồ sơ production không lưu thông tin nhạy cảm trực tiếp trong mã nguồn mà đọc DB_URL, DB_USERNAME, DB_PASSWORD, S3_BUCKET và AWS_REGION từ biến môi trường. Trên EC2, systemd quản lý tc-social, cho phép tự khởi động cùng máy chủ, khởi động lại có kiểm soát và theo dõi trạng thái dịch vụ.",
    "3.2. Luồng dữ liệu (Data Flow)": "3.2. Kiến trúc và luồng hoạt động mới",
    "Hình 8 . luồng dữ liệu": "Hình 8. Kiến trúc và luồng hoạt động mới của hệ thống trên AWS",
    "3.3. Quy trình triển khai ": "3.3. Quy trình triển khai hạ tầng và ứng dụng",
    "3.4. Luồng dữ liệu": "3.4. Luồng xử lý yêu cầu người dùng",
    "3.5. Quy trình CI/CD": "3.5. Luồng CI/CD AWS-native",
}

for p in doc.paragraphs:
    if p.text in replacements:
        set_text(p, replacements[p.text])

# Replace the numbered data-flow block.
flow = [
    "(1) Người dùng thao tác trên giao diện Thymeleaf bằng trình duyệt và gửi yêu cầu HTTP(S) đến máy chủ ứng dụng.",
    "(2) Security Group của EC2 kiểm tra kết nối đến; chỉ các cổng phục vụ ứng dụng và quản trị cần thiết được cho phép.",
    "(3) Spring Boot trên EC2 định tuyến yêu cầu qua controller, kiểm tra dữ liệu và thực thi nghiệp vụ.",
    "(4) Với dữ liệu quan hệ, tầng repository/JPA kết nối Amazon RDS for MySQL bằng các biến môi trường của hồ sơ production.",
    "(5) Với tệp và ảnh, ứng dụng sử dụng cấu hình S3_BUCKET và AWS_REGION để làm việc với Amazon S3.",
    "(6) RDS hoặc S3 trả kết quả về Spring Boot; ứng dụng dựng HTML hoặc JSON và gửi phản hồi về trình duyệt.",
]
start = next(i for i,p in enumerate(doc.paragraphs) if p.text.startswith("(1) Người dùng truy cập"))
end = next(i for i,p in enumerate(doc.paragraphs) if p.text.startswith("(9) Nginx trả"))
for i in range(start, end + 1):
    set_text(doc.paragraphs[i], flow[i-start] if i-start < len(flow) else "")

# Replace deployment steps.
steps = [
    "Bước 1: Hoàn thiện ứng dụng Spring Boot 3.2.4 trên Java 21, giao diện Thymeleaf và các module tài khoản, bài viết, tương tác, story, kết bạn, chat và quản trị.",
    "Bước 2: Khởi tạo Amazon RDS for MySQL trong mạng phù hợp; tạo database và tài khoản; giới hạn Security Group để chỉ EC2 được kết nối đến cổng 3306.",
    "Bước 3: Khởi tạo Amazon S3 cho tệp tải lên; cấu hình quyền truy cập theo nguyên tắc đặc quyền tối thiểu và gán quyền cho EC2 thông qua IAM role.",
    "Bước 4: Chuẩn bị EC2, cài Java 21 và CodeDeploy Agent; tạo dịch vụ systemd tc-social cùng thư mục /opt/tc-social.",
    "Bước 5: Cấu hình biến môi trường DB_URL, DB_USERNAME, DB_PASSWORD, S3_BUCKET và AWS_REGION cho hồ sơ production; không đưa thông tin bí mật vào kho mã.",
    "Bước 6: Tạo CodeBuild project đọc buildspec.yml để chạy mvn clean package -DskipTests và xuất JAR, appspec.yml cùng thư mục scripts.",
    "Bước 7: Tạo CodeDeploy application/deployment group cho EC2, sau đó nối nguồn mã, CodeBuild và CodeDeploy bằng CodePipeline.",
]
old_starts = ["Bước 1: Phát triển", "Bước 2: Khởi tạo", "Bước 3: Tối ưu", "Bước 4: Đóng gói", "Bước 5: Cấu hình systemd", "Bước 6: Cấu hình Nginx", "Bước 7: Cấu hình DuckDNS"]
for p in doc.paragraphs:
    for prefix, text in zip(old_starts, steps):
        if p.text.startswith(prefix): set_text(p, text)

for p in doc.paragraphs:
    if p.text == "aaaaaaaaaaaaaaaa":
        set_text(p, "Luồng xử lý được mô tả theo sáu bước ở Hình 8: trình duyệt gửi yêu cầu, EC2 tiếp nhận và Spring Boot xử lý; dữ liệu nghiệp vụ được lưu trong RDS, còn tệp được lưu trên S3; kết quả sau đó quay về trình duyệt.")
    elif p.text.startswith("Quy trình CI/CD được thực hiện bằng GitLab"):
        set_text(p, "Khi lập trình viên đẩy thay đổi lên kho mã, CodePipeline tự động kích hoạt. CodeBuild sử dụng Amazon Corretto 21 và Maven để tạo websocial-1.0-SNAPSHOT.jar. Artifact gồm JAR, appspec.yml và các script được chuyển sang CodeDeploy. CodeDeploy dừng tc-social, cài bản mới vào /opt/tc-social, cập nhật quyền sở hữu, nạp lại systemd và khởi động lại dịch vụ. Trạng thái từng giai đoạn được theo dõi tập trung trong pipeline.")
    elif p.text.strip().lower() in {"aa", "hình aa", "hình aaa", "hinh aaaaa", "thêm hình", "bảng: .ádấdsa", "bảng sádâdá", "bảng zzzzzz", "bảng aaa", "ưqewqewq"}:
        set_text(p, "")

# Synchronize the evaluation and conclusion chapters with the AWS-native architecture.
chapter_updates = {
    "Hoàn thiện quy trình triển khai ứng dụng": "Hoàn thiện quy trình triển khai tự động trên AWS bằng CodePipeline, CodeBuild, CodeDeploy và systemd, đồng thời tách dữ liệu sang RDS và S3 để nâng cao khả năng quản lý và độ tin cậy.",
    "Máy chủ Web: Nginx": "Hạ tầng ứng dụng: Amazon EC2 vận hành Spring Boot và systemd; Security Group kiểm soát các kết nối được phép.",
    "Quản lý tên miền: DuckDNS": "Dữ liệu và tệp: Amazon RDS for MySQL lưu dữ liệu quan hệ; Amazon S3 được cấu hình cho lưu trữ đối tượng.",
    "Sau khi hoàn thành quá trình thiết kế": "Sau khi hoàn thành thiết kế và phát triển, hệ thống được triển khai thử nghiệm trên AWS với Spring Boot chạy trên EC2, MySQL trên RDS và quy trình CI/CD gồm CodePipeline, CodeBuild, CodeDeploy.",
    "Nhóm đã khởi tạo thành công một máy chủ": "Nhóm đã chuẩn bị EC2 để chạy Java 21 và dịch vụ tc-social. Cơ sở dữ liệu được tách sang Amazon RDS for MySQL; thông tin kết nối được cấp qua biến môi trường production.",
    "Security Group được cấu hình": "Security Group được cấu hình theo nguyên tắc đặc quyền tối thiểu: EC2 chỉ nhận kết nối phục vụ ứng dụng và quản trị cần thiết; RDS chỉ nhận MySQL từ Security Group của EC2, không mở cổng 3306 ra Internet.",
    "Do EC2 t3.micro": "Việc tách MySQL khỏi EC2 giúp tài nguyên máy ứng dụng tập trung cho JVM và giảm ảnh hưởng lẫn nhau giữa tầng ứng dụng với tầng dữ liệu. systemd bảo đảm tc-social có thể tự khởi động và được quản lý thống nhất.",
    "Bên cạnh đó, tiến trình Java": "Hồ sơ production sử dụng biến môi trường DB_URL, DB_USERNAME, DB_PASSWORD, S3_BUCKET và AWS_REGION. Cách cấu hình này tách thông tin môi trường khỏi mã nguồn và thuận tiện khi thay đổi tài nguyên AWS.",
    "4.4. Kết quả cấu hình Nginx": "4.4. Kết quả cấu hình Amazon RDS và Amazon S3",
    "Nginx được cấu hình làm reverse proxy": "Amazon RDS for MySQL cung cấp tầng dữ liệu độc lập với EC2. Ứng dụng kết nối bằng DB_URL, DB_USERNAME và DB_PASSWORD; chỉ EC2 được phép truy cập cổng MySQL thông qua Security Group.",
    "Tên miền tcblog-cloudcomputing": "Amazon S3 được khai báo bằng S3_BUCKET và AWS_REGION để phục vụ lưu trữ tệp theo mô hình đối tượng. IAM role của EC2 nên chỉ được cấp các thao tác cần thiết trên bucket của dự án.",
    "Đối với HTTPS": "Việc tách RDS và S3 khỏi EC2 giúp dữ liệu không phụ thuộc hoàn toàn vào vòng đời của máy chủ ứng dụng, đồng thời tạo nền tảng cho sao lưu và mở rộng sau này.",
    "[Chèn ảnh truy cập": "[Chèn ảnh minh chứng RDS, S3 và Security Group sau khi cấu hình]",
    "4.5. Kết quả triển khai GitLab": "4.5. Kết quả triển khai AWS CodePipeline",
    "Sau khi đưa mã nguồn lên GitLab": "Nhóm đã bổ sung buildspec.yml và appspec.yml. CodeBuild dùng Java Corretto 21 và Maven để tạo JAR; CodeDeploy chép artifact vào /opt/tc-social và chạy các hook ApplicationStop, AfterInstall, ApplicationStart.",
    "Quá trình thực nghiệm cho thấy pipeline": "Pipeline thực hiện tuần tự Source, Build và Deploy. Khi mã nguồn thay đổi, bản JAR mới được tạo và CodeDeploy khởi động lại tc-social qua systemd, giảm thao tác sao chép thủ công và cho phép theo dõi trạng thái từng giai đoạn.",
    "[Chèn ảnh GitLab Pipeline": "[Chèn ảnh AWS CodePipeline với các giai đoạn Source, Build và Deploy thành công]",
    "Về mặt cloud, đề tài": "Về mặt cloud, đề tài kết hợp EC2, RDS, S3, IAM, Security Group và chuỗi CodePipeline–CodeBuild–CodeDeploy, qua đó thể hiện cả triển khai hạ tầng, phân tách dịch vụ và tự động hóa phát hành phần mềm.",
    "Về mặt chi phí": "Về mặt chi phí, kiến trúc mới phát sinh thêm RDS và S3 so với mô hình một EC2, nhưng đổi lại có sự tách biệt dữ liệu, quản lý vận hành rõ ràng và khả năng mở rộng tốt hơn. Cấu hình tài nguyên nhỏ phù hợp phạm vi đồ án và cần được theo dõi để tránh vượt ngân sách.",
    "Sau quá trình tìm hiểu, xây dựng": "Sau quá trình tìm hiểu, xây dựng và triển khai, nhóm đã hoàn thành ứng dụng mạng xã hội bằng Spring Boot, gồm các chức năng người dùng, bài viết, tương tác, story, kết bạn, chat và quản trị.",
    "Hệ thống có thể truy cập thông qua tên miền": "Ứng dụng được đóng gói thành JAR và vận hành trên EC2 bằng systemd; MySQL được tách sang RDS và cấu hình S3 đã được chuẩn bị cho lưu trữ tệp.",
    "Một kết quả nâng cấp quan trọng": "Kết quả nâng cấp quan trọng là quy trình CI/CD AWS-native: CodePipeline điều phối, CodeBuild đóng gói và CodeDeploy cài bản phát hành lên EC2 theo appspec.yml.",
    "Hệ thống có kiến trúc đơn giản": "Kiến trúc phân tách compute, database và object storage giúp trách nhiệm của từng thành phần rõ ràng. systemd quản lý tiến trình ứng dụng, còn CodeDeploy chuẩn hóa các bước dừng, cài và khởi động.",
    "Quy trình CI/CD bằng GitLab": "Quy trình CI/CD trên AWS giảm phụ thuộc vào truyền tệp thủ công qua SSH và cung cấp trạng thái tập trung cho từng giai đoạn Source, Build và Deploy.",
    "Việc dùng DuckDNS": "Hạn chế hiện tại là hệ thống vẫn chỉ có một EC2 nên chưa đạt độ sẵn sàng cao. Pipeline chưa có bước kiểm thử tự động bắt buộc và chiến lược rollback/blue-green; cấu hình S3 trong application-prod.properties cũng cần được hiện thực đầy đủ trong mã upload.",
    "Về CI/CD, có thể cấu hình deploy-job": "Về CI/CD, có thể bổ sung kiểm thử tự động trong CodeBuild, phê duyệt thủ công trước production, cảnh báo CloudWatch và chiến lược blue/green hoặc rollback của CodeDeploy.",
    "Với mô hình AWS EC2 t3.micro": "Với kiến trúc EC2, RDS for MySQL, S3, systemd và chuỗi CodePipeline–CodeBuild–CodeDeploy, đề tài đáp ứng mục tiêu triển khai ứng dụng web trên cloud và mô phỏng quy trình DevOps thực tế có khả năng mở rộng.",
}
for p in doc.paragraphs:
    for prefix, new_text in chapter_updates.items():
        if p.text.startswith(prefix):
            set_text(p, new_text)
            break

# Insert the new diagram into the existing figure paragraph immediately before its caption.
cap_idx = next(i for i,p in enumerate(doc.paragraphs) if p.text.strip().startswith("Hình 8"))
set_text(doc.paragraphs[cap_idx], "Hình 8. Kiến trúc và luồng hoạt động mới của hệ thống trên AWS")
img_p = doc.paragraphs[cap_idx-1]
img_p.clear()
img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_p.add_run().add_picture(str(DIAGRAM), width=Inches(6.5))
doc.paragraphs[cap_idx].alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in doc.paragraphs[cap_idx].runs:
    r.italic = True
    r.font.size = Pt(10)

# Refresh key technology and verification tables.
tech = doc.tables[3]
rows = [
    ("Front-end", "HTML, CSS, Thymeleaf", "Giao diện người dùng và quản trị."),
    ("Back-end", "Java 21, Spring Boot 3.2.4", "Nghiệp vụ, REST, xác thực và WebSocket."),
    ("Cơ sở dữ liệu", "Amazon RDS for MySQL", "Lưu dữ liệu quan hệ, tách khỏi EC2."),
    ("Lưu trữ tệp", "Amazon S3", "Lưu ảnh và tệp theo mô hình object storage."),
    ("Compute", "Amazon EC2", "Chạy tệp JAR và dịch vụ tc-social."),
    ("CI/CD", "CodePipeline, CodeBuild, CodeDeploy", "Tự động hóa từ mã nguồn đến EC2."),
    ("Vận hành", "systemd, IAM, Security Group", "Quản lý tiến trình, quyền và kết nối mạng."),
]
for i,row in enumerate(rows,1):
    for j,val in enumerate(row): tech.cell(i,j).text = val
for i in range(len(rows)+1, len(tech.rows)):
    for c in tech.rows[i].cells: c.text = ""

env = doc.tables[5]
new_env = [
    ("Compute", "Amazon EC2", "Vận hành Spring Boot"),
    ("Database", "Amazon RDS for MySQL", "Lưu dữ liệu quan hệ"),
    ("Object Storage", "Amazon S3", "Lưu tệp và ảnh"),
    ("Build", "AWS CodeBuild", "Đóng gói Maven/Java 21"),
    ("Pipeline", "AWS CodePipeline", "Điều phối các giai đoạn"),
    ("Deployment", "AWS CodeDeploy", "Triển khai artifact lên EC2"),
    ("Service Manager", "systemd", "Quản lý tc-social"),
]
for i,row in enumerate(new_env,1):
    for j,val in enumerate(row): env.cell(i,j).text = val

tests = doc.tables[4]
test_rows = [
    ("Truy cập ứng dụng", "Giao diện phản hồi từ Spring Boot trên EC2."),
    ("Kết nối RDS", "EC2 kết nối MySQL; máy ngoài không truy cập trực tiếp cổng 3306."),
    ("Chức năng người dùng", "Đăng ký, đăng nhập, hồ sơ và bảng tin hoạt động đúng."),
    ("Tương tác xã hội", "Bài viết, bình luận, thích, kết bạn và story hoạt động đúng."),
    ("Nhắn tin", "WebSocket gửi và nhận tin nhắn theo thời gian thực."),
    ("CodeBuild", "Maven tạo đúng JAR và deployment artifact."),
    ("CodeDeploy", "tc-social được cập nhật và khởi động lại thành công."),
]
for i,row in enumerate(test_rows,1):
    for j,val in enumerate(row): tests.cell(i,j).text = val

results = doc.tables[6]
result_rows = [
    ("Truy cập website", "Mở địa chỉ ứng dụng", "Giao diện phản hồi bình thường."),
    ("Đăng nhập", "Dùng tài khoản hợp lệ", "Xác thực và chuyển trang đúng."),
    ("Tương tác", "Tạo bài, thích, bình luận, kết bạn", "Dữ liệu được lưu trên RDS."),
    ("Nhắn tin", "Gửi tin qua hai phiên người dùng", "Tin nhắn hiển thị thời gian thực."),
    ("Build", "Kích hoạt CodeBuild", "JAR và artifact được tạo thành công."),
    ("Deploy", "Chạy CodeDeploy", "tc-social nhận phiên bản mới và hoạt động."),
]
for i,row in enumerate(result_rows,1):
    for j,val in enumerate(row): results.cell(i,j).text = val

# Update fields on next Word open.
settings = doc.settings._element
upd = settings.find(qn("w:updateFields"))
if upd is None:
    upd = OxmlElement("w:updateFields")
    settings.append(upd)
upd.set(qn("w:val"), "true")

doc.save(OUT)
print(OUT)
