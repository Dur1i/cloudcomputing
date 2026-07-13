from pathlib import Path

from docx import Document


SRC = Path(r"D:\Downloads-2\School\Cloud\DoAnCloud_cap_nhat.docx")
OUT = Path(r"D:\Downloads-2\School\Other\websocial\DoAnCloud_cap_nhat_2EC2.docx")


REPLACEMENTS = {
    "Trong đồ án, dù Auto Scaling Group chỉ duy trì một EC2 t3.micro để tiết kiệm chi phí, việc dùng ALB vẫn giúp thể hiện đúng mô hình triển khai cloud production hơn so với truy cập thẳng IP EC2.": (
        "Trong phiên bản hoàn thiện, Auto Scaling Group duy trì 2 EC2 t3.micro ở phía sau Application Load Balancer. "
        "Nhờ đó, ALB không chỉ đóng vai trò lớp trung gian mà còn thực hiện cân bằng tải thật sự giữa hai target healthy. "
        "Khi một EC2 hoặc service tcblog gặp lỗi, Target Group health check sẽ đánh dấu target đó unhealthy và ALB tiếp tục chuyển request sang EC2 còn hoạt động."
    ),
    "Môi trường thực nghiệm được triển khai trên AWS khu vực Asia Pacific (Singapore) - ap-southeast-1. Ứng dụng chạy trên EC2 t3.micro Ubuntu trong Auto Scaling Group với desired capacity là 1 để phù hợp với giới hạn chi phí. Spring Boot chạy bằng Java 21, được quản lý bởi systemd service tcblog.service và reverse proxy qua Nginx.": (
        "Môi trường thực nghiệm được triển khai trên AWS khu vực Asia Pacific (Singapore) - ap-southeast-1. "
        "Ứng dụng chạy trên 2 EC2 t3.micro Ubuntu trong Auto Scaling Group để chứng minh khả năng cân bằng tải và dự phòng cơ bản. "
        "Spring Boot chạy bằng Java 21, được quản lý bởi systemd service tcblog.service và reverse proxy qua Nginx trên từng EC2."
    ),
    "Hạ tầng triển khai gồm CloudFront, Application Load Balancer, Target Group, Auto Scaling Group, EC2, Nginx, Spring Boot, RDS MySQL và S3. Người dùng truy cập bằng HTTPS qua CloudFront; CloudFront chuyển request về ALB; ALB chuyển request đến EC2 healthy trong Target Group.": (
        "Hạ tầng triển khai gồm CloudFront, Application Load Balancer, Target Group, Auto Scaling Group, 2 EC2 t3.micro, Nginx, Spring Boot, RDS MySQL và S3. "
        "Người dùng truy cập bằng HTTPS qua CloudFront; CloudFront chuyển request về ALB; ALB phân phối request đến một trong hai EC2 healthy trong Target Group."
    ),
    "Target Group được cấu hình health check HTTP đến đường dẫn / hoặc /login trên port traffic 80. Khi EC2 và Nginx hoạt động đúng, target chuyển sang trạng thái healthy và ALB có thể phục vụ request từ người dùng.": (
        "Target Group được cấu hình health check HTTP đến đường dẫn / hoặc /login trên port traffic 80. "
        "Kết quả thực nghiệm cho thấy Target Group có 2 registered targets và cả 2 đều ở trạng thái Healthy. "
        "Điều này chứng minh ALB có thể điều phối request qua nhiều server, thay vì chỉ trỏ đến một EC2 duy nhất."
    ),
    "RDS giúp giảm tải cho EC2 vì MySQL không còn chạy trực tiếp trên máy chủ ứng dụng. S3 giúp tách file media khỏi ổ đĩa EC2. Khi cần tạm dừng để tiết kiệm chi phí, nhóm có thể stop EC2/RDS hoặc giảm desired capacity của Auto Scaling Group tùy giai đoạn demo.": (
        "RDS giúp giảm tải cho EC2 vì MySQL không còn chạy trực tiếp trên máy chủ ứng dụng. S3 giúp tách file media khỏi ổ đĩa EC2, nhờ đó hai EC2 có thể dùng chung dữ liệu và media. "
        "Khi cần tạm dừng để tiết kiệm chi phí, nhóm có thể giảm desired capacity của Auto Scaling Group từ 2 xuống 1 hoặc stop các tài nguyên chưa cần dùng."
    ),
    "Hệ thống đã chứng minh được ALB route đến 2 EC2 healthy. Tuy nhiên, chính sách auto scaling theo CPU/request và rolling deployment vẫn là hướng phát triển.": (
        "Hệ thống đã chứng minh được ALB route đến 2 EC2 healthy và vẫn duy trì truy cập khi một target bị stop service để demo failover. "
        "Tuy nhiên, chính sách auto scaling theo CPU/request, rolling deployment và session store dùng Redis/ElastiCache vẫn là hướng phát triển."
    ),
    "Hệ thống có khả năng mở rộng về mặt kiến trúc. Dù hiện tại Auto Scaling Group chỉ duy trì một EC2 để tiết kiệm chi phí, mô hình ALB + Target Group + ASG vẫn cho phép tăng số lượng instance trong tương lai khi cần.": (
        "Hệ thống có khả năng mở rộng về mặt kiến trúc. Auto Scaling Group hiện đã duy trì 2 EC2 healthy phía sau ALB, giúp chứng minh cân bằng tải và dự phòng cơ bản. "
        "Khi cần mở rộng hơn, có thể tăng số lượng instance, bổ sung scaling policy và phân bổ instance rõ ràng qua nhiều Availability Zone."
    ),
    "Hạn chế thứ hai là hệ thống vẫn chạy một EC2 t3.micro để kiểm soát chi phí. Điều này phù hợp cho đồ án và demo, nhưng chưa phải mô hình high availability đầy đủ vì nếu instance gặp lỗi trong thời điểm ASG chưa thay thế kịp, người dùng có thể bị gián đoạn.": (
        "Hạn chế thứ hai là hệ thống sử dụng 2 EC2 t3.micro để demo load balancing nhưng vẫn ở quy mô nhỏ nhằm kiểm soát chi phí. "
        "Mô hình này chứng minh được ALB và Target Group hoạt động, nhưng chưa phải high availability đầy đủ như production vì chưa có scaling policy hoàn chỉnh, rolling deployment và session store dùng chung."
    ),
}


APPEND_AFTER = {
    "3.2. Luồng dữ liệu và luồng người dùng": [
        "Trong cấu hình mới, ALB có 2 target healthy tương ứng với hai EC2 t3.micro. Mỗi EC2 đều chạy Nginx và Spring Boot service tcblog, cùng kết nối đến một RDS MySQL và một S3 bucket. Vì database và media đã được tách ra khỏi EC2, cả hai server có thể phục vụ cùng một ứng dụng mà không phụ thuộc vào dữ liệu cục bộ.",
    ],
    "3.4. Tổng kiến trúc hệ thống": [
        "Điểm cập nhật quan trọng của kiến trúc là Target Group hiện có 2 EC2 healthy. ALB có thể phân phối request giữa hai instance và tự động loại target unhealthy khỏi vòng phục vụ. Đây là phần chứng minh Load Balancer hoạt động thực tế trong đồ án.",
    ],
    "4.2. Kết quả triển khai hạ tầng AWS": [
        "[Cần chèn hình: Target Group hiển thị 2 registered targets, cả 2 Health status = Healthy]",
        "[Cần chèn hình: Auto Scaling Group Instance management hiển thị 2 EC2 InService]",
    ],
    "4.7. Đánh giá hệ thống": [
        "Về Load Balancer, hệ thống đã đạt được mục tiêu chứng minh cân bằng tải thật sự: Target Group có hai EC2 healthy, ALB đứng trước hai server và có thể tiếp tục phục vụ request khi một service tcblog trên một EC2 bị dừng để kiểm thử failover.",
    ],
}


def clean(text):
    return " ".join((text or "").split())


def insert_after(paragraph, text, style_name="Normal"):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_el = OxmlElement("w:p")
    paragraph._p.addnext(new_el)
    new_p = Paragraph(new_el, paragraph._parent)
    try:
        new_p.style = style_name
    except Exception:
        pass
    new_p.add_run(text)
    return new_p


def main():
    doc = Document(SRC)

    for p in doc.paragraphs:
        text = clean(p.text)
        if text in REPLACEMENTS:
            p.clear()
            p.add_run(REPLACEMENTS[text])

    for heading, additions in APPEND_AFTER.items():
        for p in doc.paragraphs:
            if clean(p.text) == heading:
                anchor = p
                for item in additions:
                    style = "Intense Quote" if item.startswith("[Cần chèn hình:") else "Normal"
                    anchor = insert_after(anchor, item, style)
                break

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
