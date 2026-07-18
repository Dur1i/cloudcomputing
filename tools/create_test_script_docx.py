from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = "TCBlog_Kich_Ban_Test_AWS.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KỊCH BẢN TEST THỰC TẾ HỆ THỐNG TC BLOG TRÊN AWS")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 77, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Phục vụ demo đồ án Cloud Computing")
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(85, 85, 85)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Kiến trúc: CloudFront -> ALB -> Auto Scaling Group/EC2 -> Nginx -> Spring Boot -> RDS/S3")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E8EEF5")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(31, 77, 120)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    set_table_width(table, [16.0])
    doc.add_paragraph()


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2


def add_steps(doc, steps):
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(step)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_test_case(doc, title, purpose, steps, expected, screenshots=None, note=None):
    add_h2(doc, title)
    add_para(doc, f"Mục đích: {purpose}")
    add_para(doc, "Các bước thực hiện:")
    add_steps(doc, steps)
    add_para(doc, "Kết quả mong đợi:")
    add_bullets(doc, expected)
    if screenshots:
        add_para(doc, "Hình ảnh nên chụp:")
        add_bullets(doc, screenshots)
    if note:
        add_callout(doc, "Ghi chú", note)


def add_summary_table(doc):
    add_h1(doc, "Bảng tổng hợp test case")
    rows = [
        ("1", "CloudFront HTTPS", "Mở URL CloudFront", "Website chạy qua HTTPS"),
        ("2", "ALB", "Tắt 1 EC2 hoặc nginx", "Web vẫn truy cập được qua target còn healthy"),
        ("3", "Auto Scaling Group", "Stop/terminate 1 EC2", "ASG duy trì số lượng instance"),
        ("4", "EC2", "Kiểm tra systemd service", "tcblog và nginx active"),
        ("5", "Nginx", "curl port 80 và port 8080", "Nginx proxy vào Spring Boot"),
        ("6", "Spring Boot", "Đăng nhập, đăng bài", "App hoạt động bình thường"),
        ("7", "RDS MySQL", "Restart app rồi kiểm tra dữ liệu", "Dữ liệu không mất"),
        ("8", "S3", "Upload ảnh bài viết/story/chat", "Object mới xuất hiện trong S3"),
        ("9", "WebSocket", "Nhắn tin bằng 2 tài khoản", "Tin nhắn realtime"),
        ("10", "SNS", "Publish message thủ công", "Email được gửi"),
        ("11", "CloudWatch Alarm", "Set alarm state hoặc tăng CPU", "Alarm gửi cảnh báo qua SNS"),
        ("12", "CloudWatch Dashboard", "Mở dashboard", "Có metric EC2/ALB/RDS"),
        ("13", "CodeBuild/SSM", "Start build", "Build succeeded và deploy lên EC2"),
        ("14", "Session Manager", "Connect vào EC2", "Quản trị EC2 không cần SSH public"),
        ("15", "Responsive UI", "Test nhiều breakpoint", "Không vỡ layout"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["STT", "Thành phần", "Cách test", "Kết quả mong đợi"]
    widths = [1.0, 3.7, 5.2, 6.1]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[idx], "1F4D78")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    set_table_width(table, widths)


def add_demo_order(doc):
    add_h1(doc, "Thứ tự demo đề xuất")
    add_steps(doc, [
        "Mở sơ đồ kiến trúc tổng thể để giới thiệu luồng User -> CloudFront -> ALB -> EC2 -> RDS/S3.",
        "Truy cập website qua CloudFront HTTPS và đăng nhập.",
        "Demo chức năng social: đăng bài, like, comment, story.",
        "Upload ảnh và mở S3 để chứng minh ảnh được lưu trên cloud storage.",
        "Demo nhắn tin realtime bằng 2 tài khoản.",
        "Mở RDS và giải thích dữ liệu user, post, comment, chat được lưu ở MySQL.",
        "Mở Target Group để chứng minh có 2 EC2 healthy.",
        "Tắt 1 EC2 hoặc nginx trên 1 EC2 để chứng minh ALB vẫn phục vụ website.",
        "Mở Auto Scaling Group để giải thích khả năng duy trì số lượng máy chủ.",
        "Mở CodeBuild và S3 releases/latest.zip để demo CI/CD.",
        "Mở Session Manager để chứng minh quản trị EC2 không cần SSH public.",
        "Publish message từ SNS để test email cảnh báo.",
        "Mở CloudWatch Dashboard và Alarm để giải thích giám sát hệ thống.",
        "Kết luận về khả năng sẵn sàng, lưu trữ cloud, giám sát và triển khai tự động.",
    ])


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    add_title(doc)
    add_callout(
        doc,
        "Mục tiêu tài liệu",
        "Tài liệu này dùng làm kịch bản kiểm thử thực tế cho đồ án TC Blog trên AWS. "
        "Trọng tâm là chứng minh hệ thống thật sự hoạt động qua các dịch vụ cloud, không chỉ dựa vào biểu đồ CloudWatch.",
    )

    add_h1(doc, "Thông tin môi trường demo")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(table.rows[0].cells[0], "Hạng mục", bold=True, color="FFFFFF")
    set_cell_text(table.rows[0].cells[1], "Giá trị demo", bold=True, color="FFFFFF")
    set_cell_shading(table.rows[0].cells[0], "1F4D78")
    set_cell_shading(table.rows[0].cells[1], "1F4D78")
    env_rows = [
        ("Region", "Asia Pacific (Singapore) - ap-southeast-1"),
        ("CloudFront URL", "https://d2tbv08v15tgw0.cloudfront.net"),
        ("ALB DNS", "tcblog-alb-1295454116.ap-southeast-1.elb.amazonaws.com"),
        ("Auto Scaling Group", "tcblog-asg, 2 EC2 t3.micro"),
        ("Target Group", "tcblog-tg"),
        ("Database", "Amazon RDS MySQL - websocial_after"),
        ("S3 Bucket", "tcblog-media-51379345"),
        ("Monitoring", "CloudWatch Dashboard, CloudWatch Alarm, SNS email alert"),
        ("CI/CD", "GitHub -> CodeBuild -> S3 artifact -> SSM Run Command -> EC2"),
    ]
    for left, right in env_rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left, bold=True)
        set_cell_text(cells[1], right)
    set_table_width(table, [4.5, 11.5])

    add_h1(doc, "Kịch bản test thực tế")

    add_test_case(
        doc,
        "1. Test truy cập website qua CloudFront HTTPS",
        "Chứng minh người dùng truy cập hệ thống qua CloudFront bằng HTTPS.",
        [
            "Mở trình duyệt và truy cập https://d2tbv08v15tgw0.cloudfront.net.",
            "Kiểm tra thanh địa chỉ có HTTPS.",
            "Đăng nhập vào hệ thống.",
            "Kiểm tra trang Home hiển thị bình thường.",
        ],
        ["Website truy cập được qua HTTPS.", "Người dùng không cần truy cập trực tiếp EC2 hoặc ALB."],
        ["Thanh địa chỉ CloudFront HTTPS.", "Trang Home sau khi đăng nhập."],
    )

    add_test_case(
        doc,
        "2. Test ALB khi tắt 1 máy chủ",
        "Chứng minh Application Load Balancer chuyển request sang EC2 còn healthy khi 1 EC2 gặp lỗi.",
        [
            "Vào EC2 -> Target Groups -> tcblog-tg.",
            "Chụp màn hình 2 target đang Healthy.",
            "Mở website qua CloudFront để xác nhận hệ thống đang chạy.",
            "Vào EC2 -> Instances và chọn 1 trong 2 EC2.",
            "Stop instance hoặc vào Session Manager rồi chạy sudo systemctl stop nginx.",
            "Quay lại Target Group, chờ 1-3 phút để target chuyển Unhealthy.",
            "Refresh website qua CloudFront nhiều lần.",
        ],
        ["Website vẫn truy cập được.", "ALB chỉ chuyển request sang EC2 còn Healthy."],
        ["Target Group trước test: 2 Healthy.", "Target Group sau test: 1 Healthy, 1 Unhealthy.", "Website vẫn truy cập được."],
        "Sau khi demo xong, bật lại máy hoặc chạy sudo systemctl start nginx và sudo systemctl start tcblog.",
    )

    add_test_case(
        doc,
        "3. Test Auto Scaling Group duy trì EC2",
        "Chứng minh ASG duy trì số lượng máy chủ theo desired capacity.",
        [
            "Vào EC2 -> Auto Scaling Groups -> tcblog-asg.",
            "Mở tab Instance management và chụp màn hình 2 EC2 InService.",
            "Stop hoặc terminate 1 EC2 trong ASG.",
            "Chờ vài phút để ASG khởi tạo hoặc duy trì instance theo desired capacity.",
        ],
        ["ASG duy trì số lượng EC2 theo cấu hình.", "Target Group nhận lại target healthy sau khi instance sẵn sàng."],
        ["Auto Scaling Group có 2 EC2 InService.", "Activity/Instance management sau khi ASG phục hồi."],
    )

    add_test_case(
        doc,
        "4. Test upload ảnh bài viết lên S3",
        "Chứng minh ảnh bài viết được lưu ở Amazon S3 thay vì lưu cục bộ trên EC2.",
        [
            "Đăng nhập website.",
            "Tạo bài viết mới và chọn ảnh.",
            "Bấm Post.",
            "Kiểm tra bài viết hiển thị ảnh.",
            "Vào S3 -> bucket tcblog-media-51379345.",
            "Tìm object mới theo thời gian upload.",
        ],
        ["Ảnh hiển thị trên bài viết.", "S3 có object mới tương ứng với ảnh vừa upload."],
        ["Bài viết có ảnh.", "Object mới trong S3 bucket."],
    )

    add_test_case(
        doc,
        "5. Test upload ảnh Story lên S3",
        "Chứng minh chức năng Story cũng sử dụng S3 để lưu media.",
        [
            "Vào Home.",
            "Bấm Add Story.",
            "Chọn ảnh và đăng Story.",
            "Kiểm tra Story xuất hiện trên giao diện.",
            "Vào S3 bucket để kiểm tra object mới.",
        ],
        ["Story tạo thành công.", "Ảnh Story được lưu trên S3."],
        ["Story trên giao diện.", "Object Story/media trong S3."],
    )

    add_test_case(
        doc,
        "6. Test nhắn tin realtime",
        "Chứng minh WebSocket hoạt động cho chức năng nhắn tin.",
        [
            "Mở trình duyệt A và đăng nhập tài khoản A.",
            "Mở trình duyệt B hoặc tab ẩn danh và đăng nhập tài khoản B.",
            "Đảm bảo 2 tài khoản có trong danh sách chat.",
            "Cả hai vào /messages.",
            "Tài khoản A gửi tin nhắn cho B.",
            "Quan sát tài khoản B không refresh trang.",
            "Gửi thử ảnh/file trong chat và kiểm tra S3.",
        ],
        ["Tin nhắn xuất hiện gần như ngay lập tức.", "File chat được upload lên S3 nếu có gửi attachment."],
        ["Hai cửa sổ trình duyệt cùng mở chat.", "Tin nhắn xuất hiện realtime.", "Object chat attachment trong S3."],
    )

    add_test_case(
        doc,
        "7. Test RDS MySQL lưu dữ liệu",
        "Chứng minh dữ liệu không phụ thuộc vào bộ nhớ EC2 mà được lưu trong Amazon RDS.",
        [
            "Đăng nhập website.",
            "Tạo bài viết mới hoặc gửi tin nhắn.",
            "Logout rồi login lại.",
            "Kiểm tra dữ liệu vẫn còn.",
            "Restart ứng dụng bằng sudo systemctl restart tcblog.",
            "Refresh website và kiểm tra dữ liệu vẫn tồn tại.",
        ],
        ["Bài viết, comment hoặc tin nhắn không mất sau khi restart app.", "Dữ liệu được lưu bền vững trong RDS MySQL."],
        ["Bài viết trước/sau restart.", "RDS database đang running."],
    )

    add_test_case(
        doc,
        "8. Test like và comment",
        "Chứng minh các tương tác mạng xã hội hoạt động và được lưu vào database.",
        [
            "Tài khoản A tạo bài viết.",
            "Tài khoản B đăng nhập.",
            "Bấm Like bài viết của A.",
            "Thêm comment.",
            "Refresh trang.",
        ],
        ["Số like và comment vẫn giữ nguyên sau refresh.", "Dữ liệu tương tác được lưu vào RDS."],
        ["Bài viết có like/comment.", "Trạng thái sau khi refresh."],
    )

    add_test_case(
        doc,
        "9. Test SNS gửi email cảnh báo",
        "Chứng minh Amazon SNS có thể gửi email cảnh báo cho người quản trị.",
        [
            "Vào AWS Console -> SNS -> Topics.",
            "Chọn topic tcblog-alerts.",
            "Kiểm tra email subscription ở trạng thái Confirmed.",
            "Bấm Publish message.",
            "Nhập Subject: Test SNS Alert - TCBlog.",
            "Nhập nội dung test và bấm Publish.",
            "Mở Gmail để kiểm tra email nhận được.",
        ],
        ["Email cảnh báo được gửi tới email đã đăng ký.", "SNS topic hoạt động độc lập với ứng dụng web."],
        ["SNS topic.", "Subscription Confirmed.", "Email nhận được."],
        "SNS dùng để gửi thông báo. Trong dự án này, SNS nhận cảnh báo từ CloudWatch Alarm và gửi email cho người quản trị.",
    )

    add_test_case(
        doc,
        "10. Test CloudWatch Alarm gửi qua SNS",
        "Chứng minh CloudWatch phát hiện tình trạng bất thường và gửi cảnh báo qua SNS.",
        [
            "Mở CloudWatch -> Alarms.",
            "Chọn alarm tcblog-ec2-high-cpu hoặc alarm phù hợp.",
            "Test thủ công bằng AWS CLI: aws cloudwatch set-alarm-state --alarm-name tcblog-ec2-high-cpu --state-value ALARM --state-reason \"Manual test alarm for demo\".",
            "Kiểm tra email SNS được gửi.",
            "Quan sát alarm tự quay về trạng thái thật theo metric sau một thời gian.",
        ],
        ["Alarm chuyển sang ALARM.", "SNS gửi email cảnh báo."],
        ["CloudWatch Alarm ở trạng thái ALARM.", "Email cảnh báo từ SNS."],
        "Có thể test bằng cách tăng CPU trên EC2, nhưng chỉ nên chạy ngắn để tránh quá tải t3.micro.",
    )

    add_test_case(
        doc,
        "11. Test CloudWatch Dashboard",
        "Chứng minh hệ thống có dashboard giám sát tập trung.",
        [
            "Vào CloudWatch -> Dashboards.",
            "Mở tcblog-dashboard.",
            "Chọn khoảng thời gian 1h hoặc 3h.",
            "Quan sát EC2 CPU, ALB Request Count, ALB 5XX, RDS CPU và RDS Database Connections.",
        ],
        ["Dashboard hiển thị metric của EC2, ALB và RDS.", "Có thể dùng để theo dõi tình trạng hệ thống khi demo."],
        ["CloudWatch Dashboard với các biểu đồ chính."],
    )

    add_test_case(
        doc,
        "12. Test CodeBuild CI/CD",
        "Chứng minh code có thể được build và deploy lên EC2 thông qua AWS.",
        [
            "Sửa một nội dung nhỏ trong project.",
            "Build local bằng mvn -q -DskipTests package.",
            "Commit và push lên GitHub.",
            "Vào CodeBuild -> tcblog-build -> Start build.",
            "Chờ trạng thái Succeeded.",
            "Vào S3 kiểm tra releases/latest.zip.",
            "Vào EC2 bằng Session Manager và chạy sudo systemctl status tcblog --no-pager.",
            "Mở website để kiểm tra thay đổi.",
        ],
        ["CodeBuild build thành công.", "Artifact latest.zip được cập nhật ở S3.", "Ứng dụng trên EC2 được deploy và chạy lại."],
        ["CodeBuild Succeeded.", "S3 releases/latest.zip.", "tcblog.service active.", "Website đã đổi nội dung."],
    )

    add_test_case(
        doc,
        "13. Test SSM Session Manager",
        "Chứng minh có thể quản trị EC2 không cần mở SSH public.",
        [
            "Vào EC2 -> Instances.",
            "Chọn một EC2 trong ASG.",
            "Bấm Connect -> Session Manager -> Connect.",
            "Chạy hostname, date, sudo systemctl status tcblog --no-pager, sudo systemctl status nginx --no-pager.",
        ],
        ["Kết nối terminal EC2 thành công.", "Không cần dùng SSH port 22 public."],
        ["Session Manager terminal.", "Service tcblog và nginx active."],
    )

    add_test_case(
        doc,
        "14. Test Nginx reverse proxy",
        "Chứng minh Nginx nhận request ở port 80 và chuyển vào Spring Boot port 8080.",
        [
            "Vào EC2 bằng Session Manager.",
            "Chạy curl -I http://127.0.0.1.",
            "Chạy curl -I http://127.0.0.1:8080/login.",
            "Chạy sudo systemctl status nginx --no-pager.",
            "Chạy sudo systemctl status tcblog --no-pager.",
            "Có thể dừng tcblog bằng sudo systemctl stop tcblog để thấy Nginx trả 502, sau đó bật lại.",
        ],
        ["Nginx active.", "Spring Boot active.", "Request port 80 được proxy vào Spring Boot."],
        ["Kết quả curl.", "Status nginx và tcblog."],
    )

    add_test_case(
        doc,
        "15. Test responsive giao diện",
        "Chứng minh giao diện dùng được trên PC, laptop, tablet và điện thoại.",
        [
            "Mở website trên desktop.",
            "Mở DevTools responsive mode.",
            "Test các kích thước 1440px, 1024px, 768px và 390px.",
            "Kiểm tra các trang Home, Explore, Messages, Profile và Login.",
            "Kiểm tra không bị tràn ngang, đè chữ, lệch nút hoặc mất chức năng.",
        ],
        ["Layout không vỡ ở nhiều kích thước.", "Các nút, form và menu mobile vẫn dùng được."],
        ["Ảnh desktop.", "Ảnh tablet.", "Ảnh mobile."],
    )

    add_summary_table(doc)
    add_demo_order(doc)

    doc.add_page_break()
    add_h1(doc, "Ghi chú chèn hình vào báo cáo")
    add_bullets(doc, [
        "CloudFront HTTPS truy cập thành công.",
        "Target Group trước/sau khi tắt 1 máy chủ.",
        "Auto Scaling Group có 2 EC2 InService.",
        "Bài viết có ảnh và object mới trong S3.",
        "Nhắn tin realtime bằng 2 tài khoản.",
        "SNS email test nhận được trong Gmail.",
        "CloudWatch Dashboard và Alarm.",
        "CodeBuild Succeeded và S3 releases/latest.zip.",
        "Session Manager kết nối EC2.",
    ])

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
