from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SRC = Path(r"D:\Downloads-2\School\Cloud\DoAnCloud.docx")
OUT = Path(r"D:\Downloads-2\School\Other\websocial\DoAnCloud_cap_nhat.docx")


def clean(text):
    return " ".join((text or "").split())


def find_para(doc, prefix):
    for p in doc.paragraphs:
        if clean(p.text).startswith(prefix):
            return p
    raise ValueError(f"Paragraph not found: {prefix}")


def set_para_text(p, text, style=None):
    p.clear()
    p.add_run(text)
    if style:
        try:
            p.style = style
        except Exception:
            pass
    return p


def insert_after(paragraph, text="", style=None, bold=False, italic=False):
    new_el = OxmlElement("w:p")
    paragraph._p.addnext(new_el)
    p = Paragraph(new_el, paragraph._parent)
    if style:
        try:
            p.style = style
        except Exception:
            pass
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    return p


def delete_paragraph(p):
    el = p._element
    parent = el.getparent()
    parent.remove(el)
    p._p = p._element = None


def replace_between(doc, start_prefix, end_prefix, lines, start_new_text=None, start_style=None):
    start = find_para(doc, start_prefix)
    if start_new_text:
        set_para_text(start, start_new_text, start_style)

    to_delete = []
    cur_el = start._p.getnext()
    while cur_el is not None:
        p = Paragraph(cur_el, start._parent)
        if clean(p.text).startswith(end_prefix):
            break
        to_delete.append(p)
        cur_el = cur_el.getnext()

    for p in to_delete:
        delete_paragraph(p)

    anchor = start
    for item in lines:
        if isinstance(item, tuple):
            text, style = item[0], item[1]
            bold = len(item) > 2 and item[2]
            italic = len(item) > 3 and item[3]
        else:
            text, style, bold, italic = item, "Normal", False, False
        anchor = insert_after(anchor, text, style, bold=bold, italic=italic)


def normal(text):
    return (text, "Normal")


def note(text):
    return (f"[Cần chèn hình: {text}]", "Intense Quote", False, True)


def main():
    doc = Document(SRC)

    replace_between(doc, "1.3.", "1.4.", [
        normal("Phạm vi chức năng của đề tài là xây dựng và triển khai một nền tảng mạng xã hội/blog mini phục vụ đăng bài, tương tác, story, hồ sơ cá nhân, nhắn tin và quản trị nội dung. Hệ thống tập trung vào luồng người dùng cơ bản gồm đăng ký, đăng nhập, xem bảng tin, đăng bài kèm hình ảnh, bình luận, thích bài viết, cập nhật thông tin cá nhân và sử dụng giao diện responsive trên máy tính, laptop, tablet và điện thoại."),
        normal("Phạm vi backend bao gồm ứng dụng Java Spring Boot, cơ chế xác thực phiên đăng nhập, xử lý nghiệp vụ mạng xã hội, kết nối cơ sở dữ liệu MySQL, lưu trữ media trên Amazon S3 và hỗ trợ các chức năng realtime bằng WebSocket cho các tương tác cần cập nhật nhanh."),
        normal("Phạm vi hạ tầng cloud được mở rộng so với bản triển khai ban đầu. Ứng dụng được đưa lên AWS theo mô hình nhiều dịch vụ: CloudFront cung cấp HTTPS public URL, Application Load Balancer phân phối request, Auto Scaling Group quản lý EC2 t3.micro chạy Ubuntu, Nginx reverse proxy vào Spring Boot, Amazon RDS MySQL lưu dữ liệu, Amazon S3 lưu ảnh và artifact, Systems Manager Parameter Store lưu cấu hình nhạy cảm, CloudWatch/SNS theo dõi và cảnh báo hệ thống."),
        normal("Phạm vi CI/CD sử dụng GitHub làm nơi quản lý mã nguồn, AWS CodeBuild để build Maven, S3 để lưu gói triển khai và AWS Systems Manager Run Command để cập nhật ứng dụng trên EC2. Do giới hạn tài khoản Free Tier, đề tài ưu tiên phương án có thể demo được, kiểm soát chi phí và có thể dừng tài nguyên khi chưa đến ngày báo cáo."),
    ])

    replace_between(doc, "1.4.", "1.5.", [
        normal("Đối tượng nghiên cứu thứ nhất là kiến trúc triển khai ứng dụng web động trên AWS, trong đó Spring Boot đóng vai trò backend, Nginx làm reverse proxy, Application Load Balancer nhận request từ Internet và CloudFront cung cấp lớp truy cập HTTPS."),
        normal("Đối tượng nghiên cứu thứ hai là cách tách các thành phần của hệ thống khỏi một máy chủ đơn lẻ: dữ liệu quan hệ được chuyển sang Amazon RDS MySQL, ảnh và file media được lưu trên Amazon S3, còn cấu hình nhạy cảm như DB_URL, DB_USERNAME, DB_PASSWORD và S3_BUCKET được quản lý bằng Parameter Store thay vì ghi cứng trong mã nguồn."),
        normal("Đối tượng nghiên cứu thứ ba là tự động hóa triển khai bằng CI/CD trên AWS. Quy trình này bao gồm GitHub, CodeBuild, S3 artifact và Systems Manager Run Command, giúp giảm thao tác thủ công khi cập nhật phiên bản ứng dụng."),
        normal("Đối tượng nghiên cứu thứ tư là vận hành và giám sát hệ thống cloud thông qua CloudWatch Dashboard, CloudWatch Alarm và SNS email alert, từ đó đánh giá trạng thái EC2, ALB, RDS và phát hiện sớm lỗi 5XX hoặc CPU cao."),
    ])

    replace_between(doc, "1.5.", "CHƯƠNG 2.", [
        normal("Ứng dụng TCBlog/TC Social được triển khai thành công trên AWS và có thể truy cập công khai qua HTTPS bằng CloudFront URL. Người dùng có thể đăng nhập, sử dụng bảng tin, đăng bài, tương tác, cập nhật hồ sơ và kiểm thử các chức năng chính trên nhiều kích thước màn hình."),
        normal("Hệ thống cloud thể hiện được kiến trúc nhiều tầng: CloudFront, ALB, EC2 Auto Scaling Group, Nginx, Spring Boot, RDS MySQL và S3. Dữ liệu bài viết, người dùng và tương tác được lưu trong RDS; ảnh bài viết, story, avatar hoặc media được lưu trong S3."),
        normal("Quy trình CI/CD được hoàn thiện ở mức phù hợp với đồ án: lập trình viên push code lên GitHub, CodeBuild build file JAR, đóng gói artifact lên S3 và gọi Systems Manager để cập nhật service trên EC2. Quy trình này giúp demo được cách tự động hóa triển khai mà không cần sử dụng CodeDeploy khi tài khoản bị giới hạn Free Tier."),
        normal("Hệ thống có dashboard và cảnh báo cơ bản trên CloudWatch/SNS để phục vụ demo vận hành: theo dõi CPU EC2, request ALB, lỗi Target 5XX, CPU RDS và kết nối database. Báo cáo cũng ghi nhận các giới hạn còn tồn tại như chưa mua domain riêng, chưa triển khai multi-instance production và cần kiểm soát chi phí khi để tài nguyên chạy lâu dài."),
    ])

    replace_between(doc, "2.3.", "2.4.", [
        normal("Một ứng dụng web hiện đại thường được chia thành nhiều lớp thay vì chỉ gồm front-end, back-end và database trên cùng một máy. Trong đề tài này, luồng xử lý bắt đầu từ trình duyệt người dùng, đi qua lớp phân phối nội dung và HTTPS, lớp cân bằng tải, máy chủ reverse proxy, ứng dụng Spring Boot, sau đó mới truy cập đến cơ sở dữ liệu và kho lưu trữ media."),
        normal("Front-end là lớp giao diện mà người dùng thao tác trực tiếp. Các trang được render bằng Thymeleaf, HTML, CSS và JavaScript, đồng thời được cải thiện để phù hợp với PC, laptop, tablet và điện thoại. Front-end chịu trách nhiệm hiển thị bảng tin, form đăng bài, story, profile, messaging, notification và các phản hồi tương tác như toast hoặc trạng thái loading."),
        normal("Back-end là lớp xử lý nghiệp vụ của hệ thống. Spring Boot tiếp nhận request, kiểm tra đăng nhập, xử lý phân quyền, tương tác với repository/service, phát sự kiện realtime qua WebSocket khi cần và trả về HTML hoặc JSON cho trình duyệt."),
        normal("Database là lớp lưu trữ dữ liệu quan hệ. Thay vì chạy MySQL trực tiếp trên EC2, hệ thống dùng Amazon RDS MySQL để tách phần dữ liệu khỏi máy chủ ứng dụng, giúp dễ quản lý, sao lưu và demo đúng tinh thần sử dụng dịch vụ cloud managed database."),
        normal("Object Storage là lớp lưu trữ file. Các ảnh bài viết, avatar, story hoặc file media được upload lên Amazon S3. Cách tách file ra khỏi EC2 giúp ứng dụng không phụ thuộc vào ổ đĩa cục bộ và phù hợp hơn với mô hình cloud."),
        note("Sơ đồ kiến trúc web app nhiều lớp: User -> CloudFront -> ALB -> EC2/Nginx -> Spring Boot -> RDS/S3"),
    ], start_new_text="2.3. Kiến trúc Web App và kiến trúc nhiều tầng trên AWS", start_style="Heading 2")

    replace_between(doc, "2.4.", "2.5.", [
        normal("Reverse Proxy là mô hình đặt một máy chủ trung gian đứng trước ứng dụng backend. Trong đề tài, Nginx chạy trên EC2 và nhận request nội bộ từ ALB, sau đó chuyển tiếp đến Spring Boot tại 127.0.0.1:8080. Cách này giúp ứng dụng Java không phải expose trực tiếp cổng 8080 ra Internet."),
        normal("Application Load Balancer là lớp tiếp nhận request ở phía AWS trước khi request đi vào EC2. ALB hoạt động cùng Target Group để kiểm tra health check của instance và chỉ chuyển request đến target còn hoạt động. Trong đồ án, dù Auto Scaling Group chỉ duy trì một EC2 t3.micro để tiết kiệm chi phí, việc dùng ALB vẫn giúp thể hiện đúng mô hình triển khai cloud production hơn so với truy cập thẳng IP EC2."),
        normal("CloudFront được sử dụng như lớp phân phối nội dung và cung cấp HTTPS public URL. Vì nhóm chưa mua domain riêng trong phạm vi Free Tier, CloudFront domain mặc định được dùng để demo HTTPS thay cho custom domain. CloudFront chuyển request về ALB origin, còn ALB/Nginx tiếp tục đưa request vào ứng dụng."),
        note("Ảnh cấu hình CloudFront origin trỏ về ALB và sơ đồ CloudFront -> ALB -> Nginx -> Spring Boot"),
    ], start_new_text="2.4. Reverse Proxy, Load Balancer và CloudFront", start_style="Heading 2")

    replace_between(doc, "2.5.", "2.6.", [
        normal("Ở bản thiết kế ban đầu, hệ thống dự kiến dùng DuckDNS để có domain động. Tuy nhiên, phiên bản triển khai mới đã chuyển sang sử dụng DNS mặc định của AWS ALB và CloudFront nhằm tận dụng dịch vụ cloud sẵn có và giảm phụ thuộc vào nhà cung cấp bên ngoài."),
        normal("CloudFront cung cấp domain dạng cloudfront.net và hỗ trợ HTTPS bằng chứng chỉ mặc định của AWS. Cách này không thay thế hoàn toàn việc dùng domain riêng như tcblog.com, nhưng phù hợp với phạm vi học tập vì không cần mua domain, không cần ACM custom certificate và vẫn chứng minh được luồng truy cập bảo mật HTTPS."),
        normal("Trong tương lai, nếu có domain riêng, hệ thống có thể mở rộng bằng Route 53 Hosted Zone, ACM certificate tại us-east-1 cho CloudFront và alias record trỏ domain về CloudFront distribution."),
        note("Ảnh CloudFront distribution domain đang truy cập HTTPS thành công"),
    ], start_new_text="2.5. DNS AWS mặc định và HTTPS với CloudFront", start_style="Heading 2")

    replace_between(doc, "2.6.", "2.7.", [
        normal("CI/CD là quy trình tự động hóa từ lúc mã nguồn thay đổi đến lúc ứng dụng được build, đóng gói và triển khai. Trong đồ án, quy trình GitLab CI/CD cũ được thay thế bằng AWS CodeBuild kết hợp GitHub, S3 và Systems Manager để phù hợp hơn với mục tiêu tận dụng dịch vụ AWS."),
        normal("GitHub đóng vai trò lưu trữ mã nguồn. CodeBuild lấy source từ repository, chạy Maven với Java 21, tạo file JAR của Spring Boot, đóng gói kèm script triển khai thành file ZIP và upload lên S3 dưới đường dẫn releases/latest.zip."),
        normal("Sau khi artifact được upload, CodeBuild gọi AWS Systems Manager Run Command đến EC2 đang nằm trong Auto Scaling Group. EC2 tải latest.zip từ S3, giải nén, copy JAR vào /opt/tcblog/websocial.jar, cập nhật systemd service tcblog.service và restart ứng dụng."),
        normal("Ưu điểm của cách triển khai này là không cần SSH trực tiếp từ máy cá nhân hoặc runner vào EC2, hạn chế phải mở cổng quản trị và vẫn phù hợp khi tài khoản Free Tier không sử dụng được CodeDeploy."),
        note("Ảnh CodeBuild build succeeded và S3 releases/latest.zip"),
    ], start_new_text="2.6. CI/CD với GitHub, CodeBuild, S3 và Systems Manager", start_style="Heading 2")

    replace_between(doc, "3.1.", "3.2.", [
        normal("Ứng dụng sử dụng Java 21 và Spring Boot làm nền tảng backend. Spring Boot phù hợp với đề tài vì hỗ trợ MVC, JPA/Hibernate, Spring Security, WebSocket và cấu trúc service/repository rõ ràng cho một hệ thống mạng xã hội nhỏ."),
        normal("Giao diện sử dụng Thymeleaf, HTML, CSS và JavaScript. Nhóm ưu tiên cải thiện trải nghiệm người dùng theo hướng hiện đại, dễ đọc, responsive và giữ logic nghiệp vụ hiện có thay vì xây dựng lại ứng dụng từ đầu."),
        normal("Cơ sở dữ liệu sử dụng Amazon RDS MySQL thay vì MySQL local trên EC2. Cách tách database ra khỏi máy chủ ứng dụng giúp mô hình gần với kiến trúc cloud thực tế hơn, đồng thời thuận tiện khi demo vai trò của managed database."),
        normal("Lưu trữ ảnh và file media sử dụng Amazon S3. Ứng dụng upload ảnh bài viết, story, avatar và media lên bucket S3, giúp giảm phụ thuộc vào filesystem của EC2."),
        normal("Hạ tầng truy cập gồm CloudFront, Application Load Balancer, Target Group, Auto Scaling Group, EC2 Ubuntu t3.micro, Nginx và systemd. Các cấu hình nhạy cảm được đặt trong Parameter Store, còn giám sát vận hành dùng CloudWatch Dashboard, CloudWatch Alarm và SNS."),
        note("Bảng công nghệ sử dụng: Spring Boot, Thymeleaf, RDS, S3, EC2, ALB, CloudFront, CodeBuild, SSM, CloudWatch"),
    ])

    replace_between(doc, "3.2.", "3.3.", [
        normal("Luồng người dùng bắt đầu khi trình duyệt truy cập CloudFront URL bằng HTTPS. CloudFront đóng vai trò edge layer, nhận request từ Internet và chuyển tiếp đến Application Load Balancer origin."),
        normal("Application Load Balancer nhận request HTTP từ CloudFront, kiểm tra Target Group và chuyển request đến EC2 còn healthy. EC2 chạy Ubuntu t3.micro, trong đó Nginx tiếp nhận request ở cổng 80 và reverse proxy vào Spring Boot tại 127.0.0.1:8080."),
        normal("Spring Boot xử lý nghiệp vụ như đăng nhập, đăng bài, story, like, comment, message, profile và admin. Khi cần dữ liệu quan hệ, ứng dụng truy vấn Amazon RDS MySQL; khi cần lưu file ảnh hoặc media, ứng dụng upload lên Amazon S3."),
        normal("Response được trả ngược lại theo luồng Spring Boot -> Nginx -> ALB -> CloudFront -> Browser. Với các chức năng realtime, WebSocket giúp cập nhật một số trạng thái tương tác nhanh hơn giữa các client."),
        note("Luồng Người Dùng: Browser -> CloudFront HTTPS -> ALB -> EC2/Nginx -> Spring Boot -> RDS/S3"),
    ], start_new_text="3.2. Luồng dữ liệu và luồng người dùng", start_style="Heading 2")

    replace_between(doc, "3.3.", "3.4.", [
        normal("Bước 1: Chuẩn bị ứng dụng Spring Boot ở local, đảm bảo build Maven thành công và cấu hình profile production đọc biến môi trường/Parameter Store thay vì hard-code thông tin database hoặc S3."),
        normal("Bước 2: Tạo EC2 t3.micro Ubuntu trong Auto Scaling Group, gán IAM Role để EC2 có thể đọc Parameter Store, truy cập S3 media/artifact và sử dụng Session Manager."),
        normal("Bước 3: Cài đặt Java 21, Maven/Git nếu cần, Nginx và cấu hình systemd service tcblog.service. Nginx proxy request từ cổng 80 vào Spring Boot cổng 8080."),
        normal("Bước 4: Tạo Amazon RDS MySQL và cập nhật chuỗi kết nối production qua Parameter Store. Ứng dụng dùng spring.datasource.url, username và password từ cấu hình cloud."),
        normal("Bước 5: Tạo S3 bucket để lưu media và artifact triển khai. Ứng dụng upload ảnh lên các prefix như posts, stories, avatars, chat hoặc uploads; CI/CD lưu gói deploy tại releases/latest.zip."),
        normal("Bước 6: Tạo Application Load Balancer, Target Group và Security Group phù hợp. Target Group health check đến EC2 qua HTTP port 80."),
        normal("Bước 7: Tạo CloudFront distribution trỏ origin về ALB, cấu hình redirect HTTP sang HTTPS và cache policy phù hợp cho ứng dụng web động."),
        normal("Bước 8: Thiết lập CodeBuild lấy source từ GitHub, build JAR, upload artifact lên S3 và gọi Systems Manager Run Command để deploy vào EC2."),
        normal("Bước 9: Tạo CloudWatch Dashboard, CloudWatch Alarm và SNS email để theo dõi EC2, ALB, RDS và nhận cảnh báo khi có lỗi."),
        note("Ảnh từng bước AWS Console: EC2/ASG, ALB Target Group healthy, RDS, S3, CloudFront, CodeBuild, CloudWatch Dashboard"),
    ], start_new_text="3.3. Quy trình triển khai trên AWS", start_style="Heading 2")

    replace_between(doc, "3.4.", "3.5.", [
        normal("Tổng kiến trúc của hệ thống gồm bốn nhóm luồng chính: luồng người dùng, luồng xử lý dữ liệu, luồng CI/CD và luồng monitoring. Các luồng này kết hợp để tạo thành mô hình vận hành cloud hoàn chỉnh cho ứng dụng TCBlog."),
        normal("Luồng người dùng đi qua CloudFront, ALB, EC2/Nginx và Spring Boot. Luồng dữ liệu tách thành hai hướng: dữ liệu quan hệ đi đến RDS MySQL, còn file media đi đến S3. Luồng CI/CD bắt đầu từ GitHub và kết thúc bằng việc restart tcblog.service trên EC2. Luồng monitoring lấy metric từ EC2, ALB, RDS và hiển thị trên CloudWatch."),
        note("Tổng Kiến Trúc: User -> CloudFront -> ALB -> ASG/EC2 -> Nginx -> Spring Boot -> RDS/S3; GitHub -> CodeBuild -> S3 -> SSM -> EC2; CloudWatch/SNS"),
    ], start_new_text="3.4. Tổng kiến trúc hệ thống", start_style="Heading 2")

    replace_between(doc, "3.5.", "3.6.", [
        normal("Quy trình CI/CD hiện tại được thiết kế theo hướng sử dụng các dịch vụ AWS có thể demo trong phạm vi Free Tier. Khi lập trình viên sửa code, mã nguồn được commit và push lên GitHub. Sau đó người quản trị có thể bấm Start build trong CodeBuild để bắt đầu quá trình build/deploy."),
        normal("Trong build phase, CodeBuild dùng môi trường Java Corretto 21, chạy mvn clean package -DskipTests để tạo file websocial-1.0-SNAPSHOT.jar. File JAR được đổi tên thành websocial.jar và đóng gói cùng thư mục deploy/scripts thành tcblog-deploy.zip."),
        normal("Trong post-build phase, CodeBuild upload tcblog-deploy.zip lên S3 theo hai key: một key có commit revision và một key cố định releases/latest.zip. Sau đó CodeBuild gọi SSM Run Command đến EC2 trong Auto Scaling Group để tải latest.zip, giải nén, copy JAR vào /opt/tcblog và restart tcblog.service."),
        normal("Trong quá trình thực hiện, nhóm cũng xử lý lỗi service name không thống nhất giữa tc-social.service và tcblog.service. Sau khi điều chỉnh script deploy, service cũ được disable và service tcblog trở thành service chính của hệ thống."),
        note("Luồng CI/CD: GitHub -> CodeBuild -> Maven JAR -> S3 releases/latest.zip -> SSM Run Command -> EC2 tcblog.service"),
    ], start_new_text="3.5. Quy trình CI/CD", start_style="Heading 2")

    replace_between(doc, "3.6.", "CHƯƠNG 4:", [
        normal("Kiểm thử hệ thống được chia thành ba nhóm: kiểm thử chức năng ứng dụng, kiểm thử hạ tầng cloud và kiểm thử CI/CD. Mục tiêu là chứng minh ứng dụng không chỉ chạy được mà còn vận hành đúng qua các dịch vụ AWS đã cấu hình."),
        normal("Kiểm thử chức năng bao gồm đăng nhập, đăng bài, đăng ảnh, story, like, comment, message, cập nhật avatar, đổi mật khẩu và kiểm tra giao diện responsive. Với chức năng upload ảnh, kết quả được xác nhận bằng việc kiểm tra object mới trong S3 bucket."),
        normal("Kiểm thử hạ tầng bao gồm truy cập CloudFront HTTPS URL, truy cập ALB URL, kiểm tra Target Group healthy, kiểm tra Nginx reverse proxy bằng curl localhost, kiểm tra tcblog.service bằng systemctl và journalctl."),
        normal("Kiểm thử CI/CD bao gồm sửa một thay đổi nhỏ trong code, commit/push lên GitHub, chạy CodeBuild, kiểm tra artifact trong S3, kiểm tra SSM command và xác nhận web đã cập nhật sau khi service restart."),
        note("Bảng test case: chức năng app, S3 upload, ALB health check, CodeBuild succeeded, CloudWatch alarm/dashboard"),
    ], start_new_text="3.6. Kiểm thử hệ thống", start_style="Heading 2")

    replace_between(doc, "CHƯƠNG 4:", "4.1.", [], start_new_text="CHƯƠNG 4. THỰC NGHIỆM VÀ ĐÁNH GIÁ", start_style="Heading 1")
    replace_between(doc, "4.1.", "4.2.", [
        normal("Môi trường thực nghiệm được triển khai trên AWS khu vực Asia Pacific (Singapore) - ap-southeast-1. Ứng dụng chạy trên EC2 t3.micro Ubuntu trong Auto Scaling Group với desired capacity là 1 để phù hợp với giới hạn chi phí. Spring Boot chạy bằng Java 21, được quản lý bởi systemd service tcblog.service và reverse proxy qua Nginx."),
        normal("Cơ sở dữ liệu sử dụng Amazon RDS MySQL với database websocial_after. File media và artifact triển khai được lưu trên S3 bucket tcblog-media-51379345. Hệ thống truy cập công khai qua CloudFront HTTPS URL, origin trỏ về Application Load Balancer."),
        note("Bảng thông tin môi trường: Region, EC2, OS, Java, ALB, RDS, S3, CloudFront, CodeBuild"),
    ], start_new_text="4.1. Môi trường thực nghiệm", start_style="Heading 2")
    replace_between(doc, "4.2.", "4.3.", [
        normal("Hạ tầng triển khai gồm CloudFront, Application Load Balancer, Target Group, Auto Scaling Group, EC2, Nginx, Spring Boot, RDS MySQL và S3. Người dùng truy cập bằng HTTPS qua CloudFront; CloudFront chuyển request về ALB; ALB chuyển request đến EC2 healthy trong Target Group."),
        normal("Target Group được cấu hình health check HTTP đến đường dẫn / hoặc /login trên port traffic 80. Khi EC2 và Nginx hoạt động đúng, target chuyển sang trạng thái healthy và ALB có thể phục vụ request từ người dùng."),
        note("Ảnh Target Group healthy, ALB DNS, CloudFront distribution enabled"),
    ], start_new_text="4.2. Kết quả triển khai hạ tầng AWS", start_style="Heading 2")
    replace_between(doc, "4.3.", "4.4.", [
        normal("Do sử dụng EC2 t3.micro, hệ thống được tối ưu để tránh vượt quá bộ nhớ. Spring Boot được chạy với tham số JVM giới hạn heap như -Xms128m và -Xmx320m, đồng thời dùng systemd để tự restart khi ứng dụng gặp lỗi."),
        normal("RDS giúp giảm tải cho EC2 vì MySQL không còn chạy trực tiếp trên máy chủ ứng dụng. S3 giúp tách file media khỏi ổ đĩa EC2. Khi cần tạm dừng để tiết kiệm chi phí, nhóm có thể stop EC2/RDS hoặc giảm desired capacity của Auto Scaling Group tùy giai đoạn demo."),
        note("Ảnh CloudWatch EC2 CPU, RDS CPU/Connections và ghi chú chi phí Free Tier"),
    ], start_new_text="4.3. Kết quả tối ưu tài nguyên và chi phí", start_style="Heading 2")
    replace_between(doc, "4.4.", "4.5.", [
        normal("Nginx được cấu hình trên EC2 để nhận request từ ALB ở cổng 80 và proxy vào Spring Boot ở 127.0.0.1:8080. Việc này giúp ẩn cổng ứng dụng Java và tạo một lớp cấu hình web server quen thuộc cho production."),
        normal("HTTPS được thực hiện ở lớp CloudFront bằng domain mặc định cloudfront.net. Do chưa mua domain riêng, hệ thống chưa sử dụng Route 53 custom domain và ACM certificate riêng, nhưng vẫn đạt mục tiêu truy cập HTTPS công khai để demo."),
        normal("CloudFront được cấu hình cache policy phù hợp cho web động, origin là ALB. Khi cập nhật giao diện, có thể dùng CloudFront invalidation với path /* để đảm bảo trình duyệt nhận bản mới."),
        note("Ảnh truy cập CloudFront HTTPS thành công và ảnh cấu hình Nginx proxy"),
    ], start_new_text="4.4. Kết quả cấu hình Nginx, ALB và CloudFront HTTPS", start_style="Heading 2")
    replace_between(doc, "4.5.", "4.6.", [
        normal("CI/CD được triển khai bằng GitHub, CodeBuild, S3 và Systems Manager. Khi build thành công, CodeBuild tạo artifact tcblog-deploy.zip, upload lên S3 releases/latest.zip và gọi SSM Run Command để deploy vào EC2."),
        normal("Quá trình thực nghiệm ghi nhận một lỗi thực tế: script deploy ban đầu cập nhật service tc-social.service trong khi service đang phục vụ web là tcblog.service. Sau khi sửa lại script, service cũ được disable và service tcblog trở thành service chính, giúp web nhận đúng bản cập nhật từ CI/CD."),
        normal("Kết quả cuối cùng cho thấy quy trình CI/CD có thể cập nhật ứng dụng mà không cần SSH thủ công. Người vận hành chỉ cần commit/push code, chạy Start build trong CodeBuild và kiểm tra trạng thái deployment qua log CodeBuild, SSM và systemctl."),
        note("Ảnh CodeBuild succeeded, S3 latest.zip, systemctl status tcblog active"),
    ], start_new_text="4.5. Kết quả triển khai CI/CD bằng CodeBuild và SSM", start_style="Heading 2")
    replace_between(doc, "4.6.", "4.7.", [
        normal("Các chức năng chính của hệ thống được kiểm thử sau khi triển khai lên cloud: đăng nhập, xem bảng tin, đăng bài, đăng ảnh, story, like, comment, nhắn tin, cập nhật hồ sơ, đổi avatar và đổi mật khẩu. Dữ liệu quan hệ được kiểm tra trong RDS, còn ảnh upload được kiểm tra trong S3 bucket."),
        normal("Kết quả kiểm thử cho thấy ứng dụng có thể hoạt động qua CloudFront/ALB thay vì chỉ chạy localhost. Các request đi qua Nginx đến Spring Boot và phản hồi về trình duyệt bình thường. Một số lỗi frontend/profile phát sinh trong quá trình cải tiến giao diện đã được sửa để modal, avatar và đổi mật khẩu hoạt động ổn định hơn."),
        note("Ảnh demo đăng bài có ảnh, S3 object mới, profile update, CloudFront URL"),
    ], start_new_text="4.6. Kết quả kiểm thử chức năng", start_style="Heading 2")
    replace_between(doc, "4.7.", "CHƯƠNG 5:", [
        normal("Về mặt chức năng, hệ thống đáp ứng yêu cầu của một nền tảng social/blog mini: người dùng có thể đăng nhập, đăng nội dung, tương tác và quản lý hồ sơ. Việc cải thiện giao diện giúp ứng dụng phù hợp hơn với nhiều thiết bị và dễ demo hơn."),
        normal("Về mặt cloud, đề tài đã mở rộng từ mô hình EC2 đơn lẻ sang kiến trúc nhiều dịch vụ AWS: CloudFront, ALB, Auto Scaling Group, EC2, Nginx, Spring Boot, RDS, S3, Parameter Store, Session Manager, CodeBuild, CloudWatch và SNS. Điều này thể hiện rõ hơn năng lực tận dụng dịch vụ cloud so với triển khai truyền thống."),
        normal("Về mặt vận hành, hệ thống có CI/CD, log systemd, Session Manager, dashboard và alarm. Tuy nhiên, do giới hạn Free Tier và chi phí, hệ thống hiện chỉ duy trì một EC2 t3.micro, chưa triển khai nhiều instance song song, chưa có domain riêng và chưa cấu hình ACM certificate cho custom domain."),
    ], start_new_text="4.7. Đánh giá hệ thống", start_style="Heading 2")

    replace_between(doc, "CHƯƠNG 5:", "5.1.", [], start_new_text="CHƯƠNG 5. KẾT LUẬN", start_style="Heading 1")
    replace_between(doc, "5.1.", "5.2.", [
        normal("Đề tài đã triển khai thành công ứng dụng TCBlog/TC Social lên AWS theo mô hình cloud nhiều tầng. Hệ thống có thể truy cập qua HTTPS bằng CloudFront URL, request được chuyển qua ALB, EC2/Nginx và Spring Boot, dữ liệu được lưu trên RDS MySQL và media được lưu trên S3."),
        normal("Nhóm đã xây dựng được quy trình CI/CD bằng GitHub, CodeBuild, S3 artifact và Systems Manager Run Command. Quy trình này giúp tự động hóa phần lớn quá trình cập nhật ứng dụng, giảm phụ thuộc vào thao tác copy file thủ công."),
        normal("Nhóm cũng hoàn thiện được các thành phần vận hành như Parameter Store, Session Manager, CloudWatch Dashboard, CloudWatch Alarm và SNS email alert để phục vụ quản trị, theo dõi và demo hệ thống."),
    ], start_new_text="5.1. Kết quả đạt được", start_style="Heading 2")
    replace_between(doc, "5.2.", "5.3.", [
        normal("Ưu điểm lớn nhất của hệ thống là thể hiện được kiến trúc cloud thực tế thay vì chỉ chạy ứng dụng trên localhost hoặc một EC2 đơn lẻ. Các dịch vụ AWS được phân chia vai trò rõ ràng: CloudFront cho HTTPS, ALB cho điều phối request, EC2 cho compute, RDS cho database, S3 cho media/artifact và CloudWatch cho monitoring."),
        normal("Hệ thống có khả năng mở rộng về mặt kiến trúc. Dù hiện tại Auto Scaling Group chỉ duy trì một EC2 để tiết kiệm chi phí, mô hình ALB + Target Group + ASG vẫn cho phép tăng số lượng instance trong tương lai khi cần."),
        normal("Quy trình triển khai có tính lặp lại và dễ demo. Khi sửa code, nhóm có thể build lại bằng CodeBuild, theo dõi log, kiểm tra artifact trên S3 và xác nhận service tcblog được restart qua SSM."),
    ], start_new_text="5.2. Ưu điểm của hệ thống", start_style="Heading 2")
    replace_between(doc, "5.3.", "5.4.", [
        normal("Hạn chế thứ nhất là hệ thống chưa có domain riêng như tcblog.com do tài khoản Free Tier không hỗ trợ mua domain trong phạm vi hiện tại. Vì vậy, HTTPS đang được demo bằng CloudFront domain mặc định thay vì custom domain."),
        normal("Hạn chế thứ hai là hệ thống vẫn chạy một EC2 t3.micro để kiểm soát chi phí. Điều này phù hợp cho đồ án và demo, nhưng chưa phải mô hình high availability đầy đủ vì nếu instance gặp lỗi trong thời điểm ASG chưa thay thế kịp, người dùng có thể bị gián đoạn."),
        normal("Hạn chế thứ ba là CI/CD hiện cần bấm Start build thủ công do webhook chưa bật hoặc chưa được dùng ổn định. Trong môi trường production, webhook hoặc pipeline hoàn chỉnh nên được cấu hình để tự động chạy khi có commit mới."),
        normal("Hạn chế thứ tư là một số chức năng realtime và giao diện vẫn còn có thể cải thiện thêm về trải nghiệm, kiểm thử tải và bảo mật nâng cao."),
    ], start_new_text="5.3. Hạn chế của hệ thống", start_style="Heading 2")
    replace_between(doc, "5.4.", "5.5.", [
        normal("Trong tương lai, hệ thống có thể mua domain riêng và cấu hình Route 53 Hosted Zone, ACM certificate, CloudFront alternate domain name để truy cập bằng URL thân thiện hơn như https://tcblog.com."),
        normal("Hệ thống có thể mở rộng Auto Scaling Group lên nhiều EC2 ở nhiều Availability Zone, cấu hình health check chặt chẽ hơn và triển khai rolling deployment để giảm downtime khi cập nhật phiên bản."),
        normal("CI/CD có thể được nâng cấp bằng webhook tự động, hoặc chuyển sang CodePipeline khi tài khoản AWS cho phép. Ngoài ra có thể thêm bước test tự động, kiểm tra bảo mật dependency và thông báo trạng thái build qua SNS hoặc email."),
        normal("Về ứng dụng, có thể tiếp tục cải thiện notification realtime, tối ưu truy vấn database, cache media, phân quyền chi tiết hơn và bổ sung kiểm thử hiệu năng."),
    ], start_new_text="5.4. Hướng phát triển trong tương lai", start_style="Heading 2")
    replace_between(doc, "5.5.", "TÀI LIỆU THAM KHẢO", [
        normal("Thông qua đề tài, nhóm đã hiểu rõ hơn sự khác biệt giữa việc chạy ứng dụng ở localhost và triển khai trên môi trường cloud. Việc kết hợp nhiều dịch vụ AWS giúp hệ thống có kiến trúc gần với thực tế hơn, đồng thời cho thấy các vấn đề cần quan tâm khi vận hành như bảo mật, chi phí, monitoring, CI/CD và khả năng mở rộng."),
        normal("Mặc dù còn một số giới hạn do phạm vi Free Tier và thời gian thực hiện, kết quả đạt được đáp ứng mục tiêu của môn học Cloud Computing: xây dựng, triển khai, vận hành và đánh giá một ứng dụng web thực tế trên nền tảng điện toán đám mây."),
    ], start_new_text="5.5. Kết luận chung", start_style="Heading 2")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
